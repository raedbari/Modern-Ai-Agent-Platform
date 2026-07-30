"""DOCX document parser.

Implements ``DocumentParser`` for ``.docx`` files using ``python-docx``.

Parsing strategy
----------------
DOCX files have no intrinsic page boundaries — page breaks are rendering
artefacts that depend on the display engine.  The content is extracted
paragraph by paragraph and the full document text is returned as a single
``ParsedPage`` with ``page_number=0``.

Paragraphs are joined with a single newline so that sentence and paragraph
boundaries are visible to the downstream chunker without introducing
spurious blank lines.  Completely empty paragraphs (common in DOCX files
as vertical spacers) are skipped.

Corruption detection
--------------------
``python-docx`` opens DOCX files as ZIP archives internally.  Corrupted
DOCX files that fail to open raise ``BadZipFile`` or ``KeyError``.  Both
are caught and re-raised as ``ParseError``.
"""

from __future__ import annotations

import io
import zipfile

import docx

from backend.app.domain.exceptions import ParseError
from backend.app.domain.ports.parser import (
    DocumentParser,
    ParsedDocument,
    ParsedPage,
    SupportedDocumentType,
)

_DOCX_MIME = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)


class DocxParser(DocumentParser):
    """Parses DOCX (``.docx``) files into a single-page ``ParsedDocument``."""

    @property
    def supported_extensions(self) -> frozenset[str]:
        return frozenset({".docx"})

    @property
    def supported_mime_types(self) -> frozenset[str]:
        return frozenset({_DOCX_MIME})

    async def parse(self, content: bytes, filename: str) -> ParsedDocument:
        """Extract paragraph text from a DOCX file.

        Args:
            content:  Raw DOCX bytes.  Must not be empty.
            filename: Original filename, used in error messages only.

        Returns:
            A ``ParsedDocument`` with a single ``ParsedPage`` containing
            all paragraph text joined by newlines.

        Raises:
            ValueError:  When ``content`` is empty.
            ParseError:  When the bytes are not a valid DOCX file or the
                         file is corrupted.
        """
        if not content:
            raise ValueError(f"Cannot parse '{filename}': file is empty.")

        try:
            document = docx.Document(io.BytesIO(content))
            paragraphs = [
                para.text
                for para in document.paragraphs
                if para.text.strip()
            ]
            text = "\n".join(paragraphs)

        except zipfile.BadZipFile as exc:
            raise ParseError(
                f"'{filename}' is not a valid DOCX file "
                f"(the file may be corrupted or is not a DOCX)."
            ) from exc
        except KeyError as exc:
            raise ParseError(
                f"'{filename}' has a malformed DOCX structure and cannot be parsed."
            ) from exc
        except Exception as exc:
            raise ParseError(
                f"An unexpected error occurred while parsing '{filename}'."
            ) from exc

        return ParsedDocument(
            document_type=SupportedDocumentType.DOCX,
            pages=(ParsedPage(page_number=0, content=text),),
        )
